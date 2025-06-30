import pdfplumber
from docx import Document
from PIL import Image
import io
import re
from typing import Dict, Any, Optional, List
from fastapi import UploadFile
from datetime import datetime
import tempfile
import os
from pathlib import Path
import asyncio
import logging

# Import our new OCR service
from .ocr_service import ocr_service, OCRLanguage, ImagePreprocessingMode

logger = logging.getLogger(__name__)

class EnhancedFileProcessor:
    """
    Enhanced file processor with intelligent content extraction,
    structure preservation, and context optimization.
    """
    
    def __init__(self):
        self.max_content_length = 12000  # Significantly increased for large files
        self.max_summary_length = 2000
        self.chunk_size = 8000  # For processing very large files in chunks
    
    async def process_file(self, file: UploadFile) -> Dict[str, Any]:
        """Process file with enhanced capabilities and return structured data."""
        
        # Extract basic file info
        file_info = {
            "filename": file.filename,
            "content_type": file.content_type,
            "size": file.size if hasattr(file, 'size') else None,
            "processed_at": datetime.now().isoformat()
        }
        
        # Process content based on file type
        if file.content_type == "application/pdf":
            result = await self._process_pdf(file)
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            result = await self._process_docx(file)
        elif file.content_type.startswith("image/"):
            result = await self._process_image(file)
        elif file.content_type.startswith("text/") or file.content_type == "text/plain":
            result = await self._process_text(file)
        else:
            # Try to process as text
            try:
                result = await self._process_text(file)
            except UnicodeDecodeError:
                result = {
                    "content": f"[Unsupported file type: {file.content_type}]",
                    "metadata": {"error": "Unsupported file type"},
                    "structure": {}
                }
        
        # Combine file info with processing results
        return {
            **file_info,
            **result,
            "optimized_content": self._optimize_content_for_context(
                result.get("content", ""),
                result.get("structure", {}),
                result.get("metadata", {})
            )
        }
    
    async def _process_pdf(self, file: UploadFile) -> Dict[str, Any]:
        """Enhanced PDF processing with structure extraction."""
        content = ""
        structure = {"pages": [], "tables": [], "headers": []}
        metadata = {}
        
        try:
            with pdfplumber.open(file.file) as pdf:
                metadata["page_count"] = len(pdf.pages)
                
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    content += f"\n--- Page {i+1} ---\n{page_text}"
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        for j, table in enumerate(tables):
                            structure["tables"].append({
                                "page": i+1,
                                "table_id": j+1,
                                "rows": len(table),
                                "cols": len(table[0]) if table else 0
                            })
                    
                    # Identify potential headers (larger text, bold, etc.)
                    headers = self._extract_headers_from_text(page_text)
                    structure["headers"].extend([
                        {"page": i+1, "text": h, "level": self._determine_header_level(h)}
                        for h in headers
                    ])
                    
                    structure["pages"].append({
                        "page": i+1,
                        "text_length": len(page_text),
                        "has_tables": bool(tables)
                    })
        
        except Exception as e:
            content = f"[Error processing PDF: {str(e)}]"
            metadata["error"] = str(e)
        
        return {
            "content": content,
            "metadata": metadata,
            "structure": structure
        }
    
    async def _process_docx(self, file: UploadFile) -> Dict[str, Any]:
        """Enhanced DOCX processing with structure extraction."""
        content = ""
        structure = {"paragraphs": [], "headers": [], "styles": []}
        metadata = {}
        
        try:
            doc = Document(file.file)
            
            # Extract core properties if available
            if hasattr(doc, 'core_properties'):
                core_props = doc.core_properties
                metadata.update({
                    "title": getattr(core_props, 'title', None),
                    "author": getattr(core_props, 'author', None),
                    "created": getattr(core_props, 'created', None),
                    "modified": getattr(core_props, 'modified', None)
                })
            
            # Process paragraphs with style information
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    style_name = paragraph.style.name if paragraph.style else "Normal"
                    
                    # Identify headers by style
                    if "Heading" in style_name or style_name.startswith("Title"):
                        level = self._extract_heading_level(style_name)
                        structure["headers"].append({
                            "text": text,
                            "level": level,
                            "style": style_name,
                            "paragraph": i
                        })
                        content += f"\n{'#' * level} {text}\n"
                    else:
                        content += f"{text}\n"
                    
                    structure["paragraphs"].append({
                        "index": i,
                        "style": style_name,
                        "length": len(text),
                        "is_header": "Heading" in style_name
                    })
            
            metadata["paragraph_count"] = len(doc.paragraphs)
            
        except Exception as e:
            content = f"[Error processing DOCX: {str(e)}]"
            metadata["error"] = str(e)
        
        return {
            "content": content,
            "metadata": metadata,
            "structure": structure
        }
    
    async def _process_image(self, file: UploadFile) -> Dict[str, Any]:
        """Enhanced image processing with OCR and metadata extraction."""
        content = ""
        metadata = {}
        structure = {}
        
        try:
            file_content = await file.read()
            img = Image.open(io.BytesIO(file_content))
            
            # Extract image metadata
            metadata.update({
                "dimensions": img.size,
                "format": img.format,
                "mode": img.mode,
                "file_size_bytes": len(file_content)
            })
            
            # Get EXIF data if available
            if hasattr(img, '_getexif') and img._getexif():
                exif_data = img._getexif()
                if exif_data:
                    metadata["exif_available"] = True
                    # Extract some useful EXIF data
                    for tag_id, value in exif_data.items():
                        if tag_id in [271, 272, 306]:  # Make, Model, DateTime
                            metadata[f"exif_tag_{tag_id}"] = str(value)
            
            # Save image to temporary file for OCR processing
            with tempfile.NamedTemporaryFile(suffix=f'.{img.format.lower()}', delete=False) as temp_file:
                img.save(temp_file.name)
                temp_path = temp_file.name
            
            try:
                # Perform OCR on the image
                ocr_result = await ocr_service.extract_text_from_image(
                    temp_path,
                    language=OCRLanguage.AUTO_DETECT,
                    preprocessing_mode=ImagePreprocessingMode.AUTO
                )
                
                # Build content with OCR results
                if ocr_result.text.strip():
                    content = f"""[Image Analysis with OCR]
Filename: {file.filename}
Dimensions: {img.size[0]}x{img.size[1]} pixels
Format: {img.format}
Color Mode: {img.mode}
File Size: {len(file_content):,} bytes

=== EXTRACTED TEXT ===
{ocr_result.text}

=== OCR METADATA ===
Confidence: {ocr_result.confidence:.1f}%
Detected Language: {ocr_result.language}
Processing Time: {ocr_result.processing_time:.2f}s
Word Count: {ocr_result.word_count}
Character Count: {ocr_result.character_count}
Preprocessing Applied: {', '.join(ocr_result.preprocessing_applied)}"""
                    
                    # Add OCR metadata
                    metadata.update({
                        "ocr_confidence": ocr_result.confidence,
                        "ocr_language": ocr_result.language,
                        "ocr_word_count": ocr_result.word_count,
                        "ocr_character_count": ocr_result.character_count,
                        "ocr_processing_time": ocr_result.processing_time,
                        "ocr_preprocessing": ocr_result.preprocessing_applied,
                        "ocr_errors": ocr_result.errors
                    })
                    
                    # Add structure information
                    structure.update({
                        "has_text": True,
                        "text_quality": "high" if ocr_result.confidence > 80 else "medium" if ocr_result.confidence > 60 else "low",
                        "extracted_text": ocr_result.text
                    })
                else:
                    content = f"""[Image Analysis - No Text Detected]
Filename: {file.filename}
Dimensions: {img.size[0]}x{img.size[1]} pixels
Format: {img.format}
Color Mode: {img.mode}
File Size: {len(file_content):,} bytes

=== OCR RESULTS ===
No readable text was detected in this image.
OCR Confidence: {ocr_result.confidence:.1f}%
Processing Time: {ocr_result.processing_time:.2f}s
Preprocessing Applied: {', '.join(ocr_result.preprocessing_applied)}"""
                    
                    metadata.update({
                        "ocr_confidence": ocr_result.confidence,
                        "ocr_no_text_detected": True,
                        "ocr_processing_time": ocr_result.processing_time
                    })
                    
                    structure["has_text"] = False
            
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_path)
                except Exception:
                    pass
            
        except Exception as e:
            logger.error(f"Error processing image {file.filename}: {str(e)}")
            content = f"[Error processing image: {str(e)}]"
            metadata["error"] = str(e)
        
        return {
            "content": content,
            "metadata": metadata,
            "structure": structure
        }
    
    async def _process_text(self, file: UploadFile) -> Dict[str, Any]:
        """Enhanced text processing with structure analysis."""
        content = ""
        metadata = {}
        structure = {"lines": [], "sections": []}
        
        try:
            file_content = await file.read()
            content = file_content.decode("utf-8")
            
            lines = content.split('\n')
            metadata.update({
                "line_count": len(lines),
                "character_count": len(content),
                "word_count": len(content.split())
            })
            
            # Analyze structure
            for i, line in enumerate(lines):
                line_info = {
                    "line_number": i + 1,
                    "length": len(line),
                    "is_empty": not line.strip()
                }
                
                # Detect potential headers/sections
                if line.strip():
                    if line.isupper() or line.startswith('#') or self._looks_like_header(line):
                        line_info["is_header"] = True
                        structure["sections"].append({
                            "title": line.strip(),
                            "line": i + 1
                        })
                
                structure["lines"].append(line_info)
            
        except UnicodeDecodeError as e:
            content = f"[Error: Unable to decode text file - {str(e)}]"
            metadata["error"] = "encoding_error"
        except Exception as e:
            content = f"[Error processing text file: {str(e)}]"
            metadata["error"] = str(e)
        
        return {
            "content": content,
            "metadata": metadata,
            "structure": structure
        }
    
    def _optimize_content_for_context(self, content: str, structure: Dict, metadata: Dict) -> str:
        """Optimize content for AI context with intelligent chunking and summarization."""
        if not content:
            return content
            
        # For small files, return as-is
        if len(content) <= self.max_content_length:
            return content
        
        # Create optimized version with metadata summary
        optimization_summary = self._create_metadata_summary(metadata, structure)
        
        # For very large files, use intelligent chunking strategy
        if len(content) > self.max_content_length * 3:  # 36K+ chars
            return self._create_large_file_summary(content, structure, metadata, optimization_summary)
        
        # For moderately large files, use enhanced smart truncation
        available_space = self.max_content_length - len(optimization_summary) - 200
        
        if available_space <= 0:
            return optimization_summary + "\n[Content too large to include]"
        
        # Enhanced smart truncation - prioritize important sections
        optimized_content = self._enhanced_smart_truncate(content, available_space, structure)
        
        return f"{optimization_summary}\n\n{optimized_content}"
    
    def _create_metadata_summary(self, metadata: Dict, structure: Dict) -> str:
        """Create a concise metadata summary."""
        summary_parts = []
        
        if metadata.get("page_count"):
            summary_parts.append(f"{metadata['page_count']} pages")
        if metadata.get("word_count"):
            summary_parts.append(f"{metadata['word_count']} words")
        if metadata.get("title"):
            summary_parts.append(f"Title: {metadata['title']}")
        if metadata.get("author"):
            summary_parts.append(f"Author: {metadata['author']}")
        
        # Structure info
        if structure.get("headers"):
            header_count = len(structure["headers"])
            summary_parts.append(f"{header_count} sections")
        if structure.get("tables"):
            table_count = len(structure["tables"])
            summary_parts.append(f"{table_count} tables")
        
        return f"[Document: {', '.join(summary_parts)}]" if summary_parts else "[Document processed]"
    
    def _smart_truncate(self, content: str, max_length: int, structure: Dict) -> str:
        """Intelligently truncate content preserving important sections."""
        if len(content) <= max_length:
            return content
        
        # If we have headers, try to include complete sections
        headers = structure.get("headers", [])
        if headers:
            return self._truncate_by_sections(content, max_length, headers)
        
        # Fallback to smart paragraph truncation
        return self._truncate_by_paragraphs(content, max_length)
    
    def _truncate_by_sections(self, content: str, max_length: int, headers: List[Dict]) -> str:
        """Truncate content by preserving complete sections."""
        # Try to fit the first few complete sections
        sections = content.split('\n')
        result = []
        current_length = 0
        
        for line in sections[:50]:  # Limit to first 50 lines for efficiency
            if current_length + len(line) + 1 > max_length:
                break
            result.append(line)
            current_length += len(line) + 1
        
        if current_length < len(content):
            result.append("\n[Content truncated - key sections preserved]")
        
        return '\n'.join(result)
    
    def _truncate_by_paragraphs(self, content: str, max_length: int) -> str:
        """Truncate content by preserving complete paragraphs."""
        paragraphs = content.split('\n\n')
        result = []
        current_length = 0
        
        for para in paragraphs:
            if current_length + len(para) + 2 > max_length:
                break
            result.append(para)
            current_length += len(para) + 2
        
        if current_length < len(content):
            result.append("\n[Content truncated - complete paragraphs preserved]")
        
        return '\n\n'.join(result)
    
    def _extract_headers_from_text(self, text: str) -> List[str]:
        """Extract potential headers from text using heuristics."""
        lines = text.split('\n')
        headers = []
        
        for line in lines:
            line = line.strip()
            if line and (line.isupper() or self._looks_like_header(line)):
                headers.append(line)
        
        return headers
    
    def _looks_like_header(self, line: str) -> bool:
        """Determine if a line looks like a header using heuristics."""
        line = line.strip()
        if not line:
            return False
        
        # Check for common header patterns
        header_patterns = [
            r'^\d+\.\s+[A-Z]',  # "1. Introduction"
            r'^[A-Z][A-Z\s]+$',  # All caps
            r'^#+\s',  # Markdown headers
            r'^[A-Z][^.]*:$',  # Ends with colon
        ]
        
        return any(re.match(pattern, line) for pattern in header_patterns)
    
    def _determine_header_level(self, header: str) -> int:
        """Determine header level based on formatting."""
        if header.startswith('###'):
            return 3
        elif header.startswith('##'):
            return 2
        elif header.startswith('#'):
            return 1
        elif header.isupper():
            return 1
        else:
            return 2
    
    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from Word style name."""
        if "Heading" in style_name:
            try:
                return int(re.search(r'\d+', style_name).group())
            except (AttributeError, ValueError):
                return 1
        elif style_name.startswith("Title"):
            return 1
        else:
            return 2

    def _create_large_file_summary(self, content: str, structure: Dict, metadata: Dict, meta_summary: str) -> str:
        """Create an intelligent summary for very large files."""
        summary_parts = [meta_summary]
        
        # Extract key sections based on structure
        if structure.get("headers"):
            summary_parts.append("\n=== DOCUMENT STRUCTURE ===")
            headers = structure["headers"][:20]  # Limit to first 20 headers
            for header in headers:
                level_prefix = "  " * (header.get("level", 1) - 1)
                summary_parts.append(f"{level_prefix}- {header.get('text', '')}")
        
        # Add beginning of document
        summary_parts.append("\n=== DOCUMENT BEGINNING ===")
        beginning = content[:2000]
        summary_parts.append(beginning)
        
        # Try to extract key sections from middle
        if len(content) > 4000:
            summary_parts.append("\n=== KEY SECTIONS ===")
            middle_sections = self._extract_key_sections(content, structure)
            summary_parts.extend(middle_sections)
        
        # Add document ending
        if len(content) > 6000:
            summary_parts.append("\n=== DOCUMENT ENDING ===")
            ending = content[-1500:]
            summary_parts.append(ending)
        
        summary_parts.append(f"\n[Note: This is an intelligent summary of a {len(content):,} character document. Key sections and structure preserved.]")
        
        result = "\n".join(summary_parts)
        
        # Ensure we don't exceed limits
        if len(result) > self.max_content_length:
            result = result[:self.max_content_length - 100] + "\n[Summary truncated]"
        
        return result
    
    def _extract_key_sections(self, content: str, structure: Dict) -> List[str]:
        """Extract key sections from the middle of large documents."""
        sections = []
        lines = content.split('\n')
        total_lines = len(lines)
        
        # Look for important keywords or patterns
        important_keywords = [
            'conclusion', 'summary', 'results', 'findings', 'recommendations',
            'abstract', 'executive summary', 'key points', 'highlights',
            'introduction', 'methodology', 'analysis', 'discussion'
        ]
        
        # Scan through document for important sections
        current_section = []
        section_count = 0
        max_sections = 3
        
        for i, line in enumerate(lines[100:-100]):  # Skip beginning and end
            line_lower = line.lower().strip()
            
            # Check if this line indicates an important section
            is_important = any(keyword in line_lower for keyword in important_keywords)
            is_header = self._looks_like_header(line) or line.isupper()
            
            if is_important or is_header:
                # Save previous section if it exists
                if current_section and section_count < max_sections:
                    section_text = '\n'.join(current_section[:15])  # Limit section length
                    sections.append(f"\n--- Section {section_count + 1} ---\n{section_text}")
                    section_count += 1
                    current_section = []
                
                if section_count >= max_sections:
                    break
            
            current_section.append(line)
            
            # Limit section size
            if len(current_section) > 20:
                current_section = current_section[-15:]  # Keep last 15 lines
        
        # Add last section if we have room
        if current_section and section_count < max_sections:
            section_text = '\n'.join(current_section[:15])
            sections.append(f"\n--- Section {section_count + 1} ---\n{section_text}")
        
        return sections
    
    def _enhanced_smart_truncate(self, content: str, max_length: int, structure: Dict) -> str:
        """Enhanced smart truncation with better section preservation."""
        if len(content) <= max_length:
            return content
        
        # Strategy 1: If we have clear structure, preserve complete sections
        headers = structure.get("headers", [])
        if headers and len(headers) > 2:
            return self._truncate_by_complete_sections(content, max_length, headers)
        
        # Strategy 2: Preserve beginning + important middle + end
        return self._truncate_with_sandwich_approach(content, max_length)
    
    def _truncate_by_complete_sections(self, content: str, max_length: int, headers: List[Dict]) -> str:
        """Truncate by keeping complete sections."""
        lines = content.split('\n')
        result = []
        current_length = 0
        sections_included = 0
        max_sections = min(5, len(headers))  # Include up to 5 sections
        
        # Include document beginning
        beginning_lines = lines[:20]
        for line in beginning_lines:
            if current_length + len(line) + 1 > max_length * 0.4:  # Use 40% for beginning
                break
            result.append(line)
            current_length += len(line) + 1
        
        # Try to include some complete sections
        section_start_indicators = [h.get('text', '') for h in headers[:max_sections]]
        
        i = len(result)
        while i < len(lines) and sections_included < max_sections:
            line = lines[i]
            if any(indicator in line for indicator in section_start_indicators if indicator):
                # Start of a new section - try to include it completely
                section_lines = [line]
                j = i + 1
                
                # Collect lines until next section or reasonable limit
                while j < len(lines) and len(section_lines) < 50:
                    next_line = lines[j]
                    if any(indicator in next_line for indicator in section_start_indicators[sections_included+1:] if indicator):
                        break  # Found next section
                    section_lines.append(next_line)
                    j += 1
                
                # Check if we can fit this section
                section_text = '\n'.join(section_lines)
                if current_length + len(section_text) <= max_length * 0.9:  # Leave 10% buffer
                    result.extend(section_lines)
                    current_length += len(section_text)
                    sections_included += 1
                    i = j
                else:
                    break
            else:
                i += 1
        
        if current_length < len(content):
            result.append(f"\n[Document truncated - {sections_included} of {len(headers)} sections included]")
        
        return '\n'.join(result)
    
    def _truncate_with_sandwich_approach(self, content: str, max_length: int) -> str:
        """Use beginning + middle + end approach for documents without clear structure."""
        # Allocate space: 50% beginning, 25% middle, 25% end
        beginning_length = int(max_length * 0.5)
        middle_length = int(max_length * 0.25)
        end_length = int(max_length * 0.25)
        
        beginning = content[:beginning_length]
        end = content[-end_length:]
        
        # Find interesting middle section
        content_length = len(content)
        middle_start = int(content_length * 0.4)  # Start at 40% through document
        middle_end = middle_start + middle_length * 3  # Sample 3x the space we need
        middle_sample = content[middle_start:middle_end]
        
        # Try to find a coherent section in the middle
        middle_paragraphs = middle_sample.split('\n\n')
        middle_content = []
        middle_current_length = 0
        
        for para in middle_paragraphs:
            if middle_current_length + len(para) + 2 > middle_length:
                break
            middle_content.append(para)
            middle_current_length += len(para) + 2
        
        middle = '\n\n'.join(middle_content)
        
        return f"{beginning}\n\n[... middle sections omitted ...]\n\n{middle}\n\n[... sections omitted ...]\n\n{end}\n\n[Document truncated - key sections preserved]"

# Create global instance
file_processor = EnhancedFileProcessor()

# Maintain backward compatibility
async def process_file(file: UploadFile):
    """Backward compatible function that returns just content."""
    result = await file_processor.process_file(file)
    return result.get("optimized_content", result.get("content", ""))

